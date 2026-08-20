from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "afgri-partnership-content-template.docx"

LOGO = ROOT / "public" / "Logo.png"
AFGRI_LOGO = ROOT / "public" / "afgri-logo-colour-cropped.png"
ATLAS_LOGO = ROOT / "public" / "atlas" / "atlas-logo-horizontal-dark.png"

INK = RGBColor(0, 29, 46)
GRAPHITE = RGBColor(43, 60, 72)
MUTED = RGBColor(99, 115, 126)
BLUE = RGBColor(0, 67, 243)
PALE_BLUE_RGB = RGBColor(193, 217, 229)
PALE_BLUE = "C1D9E5"
SOFT_BLUE = "EEF6FA"
PALE_GRAY = "F7FAFC"
LINE = "C1D9E5"
NAVY = "001D2E"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, widths, align=WD_TABLE_ALIGNMENT.CENTER):
    table.alignment = align
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def clear_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text(paragraph, text, size=10.5, color=INK, bold=False, italic=False, name="Arial"):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic, name=name)
    return run


def style_paragraph(paragraph, before=0, after=8, line=1.25, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run, size=8.2, color=MUTED)


def blank_line(doc, pts):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    return p


def setup_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)
    section.different_first_page_header_footer = False

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        s = doc.styles[style_name]
        s.font.name = "Arial"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        s.font.size = Pt(9.8)
        s.paragraph_format.space_after = Pt(3)
        s.paragraph_format.line_spacing = 1.15


def set_running_furniture(section):
    header = section.header
    for element in list(header._element):
        if element.tag in (qn("w:p"), qn("w:tbl")):
            header._element.remove(element)
    table = header.add_table(rows=1, cols=3, width=Inches(6.94))
    set_table_fixed_width(table, [2100, 4660, 3234])
    clear_table_borders(table)
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_margins(cell, top=0, bottom=55, start=0, end=0)
        set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LINE})
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.0)
        p.alignment = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT][idx]
    add_text(table.cell(0, 0).paragraphs[0], "Smart Steel x Afgri", size=8.2, color=MUTED, bold=True)
    add_text(table.cell(0, 1).paragraphs[0], "Pilot Partnership Framework", size=8.2, color=MUTED)
    add_text(table.cell(0, 2).paragraphs[0], "Draft for discussion", size=8.2, color=MUTED)

    footer = section.footer
    for element in list(footer._element):
        if element.tag in (qn("w:p"), qn("w:tbl")):
            footer._element.remove(element)
    ft = footer.add_table(rows=1, cols=3, width=Inches(6.94))
    set_table_fixed_width(ft, [3900, 3094, 3000])
    clear_table_borders(ft)
    for idx, cell in enumerate(ft.rows[0].cells):
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
        p = cell.paragraphs[0]
        p.alignment = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT][idx]
        style_paragraph(p, after=0, line=1.0)
    add_text(ft.cell(0, 0).paragraphs[0], "Subject to final agreement", size=8.0, color=MUTED)
    add_text(ft.cell(0, 1).paragraphs[0], "Atlas-enabled partnership workflow", size=8.0, color=MUTED)
    p = ft.cell(0, 2).paragraphs[0]
    add_text(p, "Page ", size=8.0, color=MUTED)
    add_page_number(p)


def add_cover_rule(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_fixed_width(table, [7700, 2294])
    clear_table_borders(table)
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
        if idx == 0:
            set_cell_border(cell, top={"val": "single", "sz": "10", "color": NAVY})
        else:
            set_cell_shading(cell, NAVY)
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.0)
            add_text(p, " ", size=4)


def add_cover_logos(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_fixed_width(table, [2500, 4994, 2500])
    clear_table_borders(table)
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_margins(cell, top=130, bottom=0, start=0, end=0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.0)
        p.alignment = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT][idx]
    table.cell(0, 0).paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(1.25))
    add_text(table.cell(0, 1).paragraphs[0], "+", size=22, color=BLUE, bold=True)
    table.cell(0, 2).paragraphs[0].add_run().add_picture(str(AFGRI_LOGO), width=Inches(1.35))


def add_cover_meta(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_fixed_width(table, [2350, 2940, 2350, 2354])
    clear_table_borders(table)
    labels = [
        ("Document", "Discussion paper"),
        ("Status", "Draft for alignment"),
        ("Pilot focus", "Afgri-led market activation"),
        ("Platform", "Atlas-enabled workflow"),
    ]
    for idx, (label, value) in enumerate(labels):
        cell = table.cell(0, idx)
        set_cell_shading(cell, PALE_BLUE if idx in (0, 2) else "FFFFFF")
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": LINE},
            bottom={"val": "single", "sz": "4", "color": LINE},
            left={"val": "single", "sz": "4", "color": LINE},
            right={"val": "single", "sz": "4", "color": LINE},
        )
        set_cell_margins(cell, top=150, bottom=160, start=160, end=140)
        p = cell.paragraphs[0]
        style_paragraph(p, after=3, line=1.0)
        add_text(p, label.upper(), size=7.5, color=BLUE, bold=True)
        p2 = cell.add_paragraph()
        style_paragraph(p2, after=0, line=1.08)
        add_text(p2, value, size=9.3, color=INK, bold=idx == 1)


def add_cover(doc):
    first_footer = doc.sections[0].first_page_footer
    p = first_footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, after=0, line=1.0)
    add_text(p, "Prepared by Smart Steel | Subject to final agreement", size=8.0, color=MUTED)

    add_cover_rule(doc)
    add_cover_logos(doc)
    blank_line(doc, 106)

    marker = doc.add_paragraph()
    style_paragraph(marker, after=0, line=1.0)
    add_text(marker, "+", size=44, color=BLUE, bold=True)

    p = doc.add_paragraph()
    style_paragraph(p, after=0, line=0.92)
    add_text(p, "Pilot", size=56, color=INK, bold=True)
    p = doc.add_paragraph()
    style_paragraph(p, after=0, line=0.92)
    add_text(p, "Partnership", size=54, color=INK, bold=True)
    p = doc.add_paragraph()
    style_paragraph(p, after=20, line=0.92)
    add_text(p, "Framework", size=50, color=INK, bold=True)

    sub = doc.add_paragraph()
    style_paragraph(sub, after=8, line=1.12)
    add_text(sub, "Smart Steel x Afgri", size=15, color=BLUE, bold=True)
    sub2 = doc.add_paragraph()
    style_paragraph(sub2, after=96, line=1.22)
    add_text(
        sub2,
        "A collaborative discussion paper for a controlled pilot across Atlas Warehouses, Solar Ground Mounts, Solar Carports and request-only agricultural structures.",
        size=10.8,
        color=MUTED,
    )

    add_cover_meta(doc)


def add_section_opener(doc, number, title, subtitle, cards, start_new_page=True):
    if start_new_page:
        doc.add_page_break()
    grid = doc.add_table(rows=1, cols=2)
    set_table_fixed_width(grid, [1320, 8674])
    clear_table_borders(grid)
    ncell, tcell = grid.cell(0, 0), grid.cell(0, 1)
    set_cell_shading(ncell, PALE_BLUE)
    set_cell_margins(ncell, top=155, bottom=155, start=130, end=130)
    set_cell_margins(tcell, top=12, bottom=0, start=360, end=0)
    for cell in [ncell, tcell]:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = ncell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, after=0, line=1.0)
    add_text(p, number, size=22, color=BLUE, bold=True)

    p = tcell.paragraphs[0]
    style_paragraph(p, after=2, line=0.96)
    add_text(p, title, size=31, color=INK, bold=False)
    p2 = tcell.add_paragraph()
    style_paragraph(p2, after=0, line=1.12)
    add_text(p2, subtitle, size=9.4, color=MUTED)

    blank_line(doc, 18)
    lead_table = doc.add_table(rows=1, cols=2)
    set_table_fixed_width(lead_table, [220, 7560], align=WD_TABLE_ALIGNMENT.LEFT)
    clear_table_borders(lead_table)
    rail, lead_cell = lead_table.cell(0, 0), lead_table.cell(0, 1)
    set_cell_shading(rail, "0043F3")
    set_cell_margins(rail, top=0, bottom=0, start=0, end=0)
    set_cell_margins(lead_cell, top=0, bottom=0, start=210, end=0)
    lead = lead_cell.paragraphs[0]
    style_paragraph(lead, after=0, line=1.24)
    add_text(
        lead,
        "A collaborative discussion paper for a controlled pilot between Afgri and Smart Steel, built around clear commercial logic, practical activation and a measured path to rollout.",
        size=10.5,
        color=GRAPHITE,
    )
    blank_line(doc, 22)

    ctable = doc.add_table(rows=1, cols=3)
    set_table_fixed_width(ctable, [3160, 3160, 3674])
    clear_table_borders(ctable)
    for idx, (label, text, fill) in enumerate(cards):
        cell = ctable.cell(0, idx)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=215, bottom=215, start=210, end=210)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": LINE},
            bottom={"val": "single", "sz": "4", "color": LINE},
            left={"val": "single", "sz": "4", "color": LINE},
            right={"val": "single", "sz": "4", "color": LINE},
        )
        p = cell.paragraphs[0]
        style_paragraph(p, after=10, line=1.0)
        add_text(p, label.upper(), size=7.2, color=BLUE, bold=True)
        p2 = cell.add_paragraph()
        style_paragraph(p2, after=0, line=1.22)
        add_text(p2, text, size=8.8, color=INK)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        style_paragraph(p, before=16, after=8, line=1.0)
        add_text(p, text, size=18, color=INK, bold=True)
    else:
        style_paragraph(p, before=14, after=6, line=1.08)
        add_text(p, text, size=13.2, color=GRAPHITE, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, after=7, line=1.26)
    add_text(p, text, size=10.3, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, after=3, line=1.15)
    if p.runs:
        p.runs[0].text = ""
    add_text(p, text, size=9.7, color=INK)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=2)
    set_table_fixed_width(table, [820, 9174])
    clear_table_borders(table)
    left, right = table.cell(0, 0), table.cell(0, 1)
    set_cell_shading(left, PALE_BLUE)
    set_cell_shading(right, PALE_GRAY)
    for cell in [left, right]:
        set_cell_margins(cell, top=150, bottom=150, start=150, end=150)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": LINE},
            bottom={"val": "single", "sz": "4", "color": LINE},
            left={"val": "single", "sz": "4", "color": LINE},
            right={"val": "single", "sz": "4", "color": LINE},
        )
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, after=0, line=1.0)
    add_text(p, "+", size=24, color=BLUE, bold=True)
    p = right.paragraphs[0]
    style_paragraph(p, after=4, line=1.05)
    add_text(p, label.upper(), size=7.8, color=BLUE, bold=True)
    p2 = right.add_paragraph()
    style_paragraph(p2, after=0, line=1.22)
    add_text(p2, text, size=9.8, color=INK)
    blank_line(doc, 5)


def add_sample_table(doc):
    table = doc.add_table(rows=4, cols=4)
    set_table_fixed_width(table, [1700, 3500, 2994, 1800])
    clear_table_borders(table)
    headers = ["Area", "Position", "Owner", "Status"]
    for c, text in enumerate(headers):
        cell = table.cell(0, c)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=120, bottom=120, start=120, end=120)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.0)
        add_text(p, text, size=8.8, color=RGBColor(255, 255, 255), bold=True)
    rows = [
        ("Commercial", "Protected dealer-price model with 5% pilot return.", "Joint", "Align"),
        ("Marketing", "Afgri-led campaign with Smart Steel support material.", "Afgri", "Align"),
        ("Atlas", "Lead capture, product workflow and quote status.", "Smart Steel", "Pilot"),
    ]
    for r, row in enumerate(rows, start=1):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            set_cell_shading(cell, "FFFFFF")
            set_cell_margins(cell, top=120, bottom=120, start=120, end=120)
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": "4", "color": LINE},
                left={"val": "single", "sz": "4", "color": LINE},
                right={"val": "single", "sz": "4", "color": LINE},
            )
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.12)
            add_text(p, text, size=8.8, color=INK if c != 3 else BLUE, bold=c == 0)


def add_topline(doc, kicker, title, subtitle=None):
    table = doc.add_table(rows=1, cols=2)
    set_table_fixed_width(table, [1420, 8574])
    clear_table_borders(table)
    left, right = table.cell(0, 0), table.cell(0, 1)
    set_cell_shading(left, PALE_BLUE)
    set_cell_margins(left, top=105, bottom=105, start=80, end=80)
    set_cell_margins(right, top=0, bottom=0, start=185, end=0)
    for cell in [left, right]:
        set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LINE})
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, after=0, line=1.0)
    add_text(p, kicker, size=8.2, color=BLUE, bold=True)
    p = right.paragraphs[0]
    style_paragraph(p, after=2, line=1.0)
    add_text(p, title, size=18, color=INK, bold=True)
    if subtitle:
        p2 = right.add_paragraph()
        style_paragraph(p2, after=0, line=1.12)
        add_text(p2, subtitle, size=9.4, color=MUTED)
    blank_line(doc, 18)


def add_signal_cards(doc, cards):
    table = doc.add_table(rows=1, cols=len(cards))
    widths = [int(9994 / len(cards))] * len(cards)
    widths[-1] += 9994 - sum(widths)
    set_table_fixed_width(table, widths)
    clear_table_borders(table)
    for idx, card in enumerate(cards):
        label, title, body, fill = card
        cell = table.cell(0, idx)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=145, bottom=155, start=150, end=150)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": LINE},
            bottom={"val": "single", "sz": "4", "color": LINE},
            left={"val": "single", "sz": "4", "color": LINE},
            right={"val": "single", "sz": "4", "color": LINE},
        )
        p = cell.paragraphs[0]
        style_paragraph(p, after=7, line=1.0)
        add_text(p, label.upper(), size=7.2, color=BLUE, bold=True)
        p2 = cell.add_paragraph()
        style_paragraph(p2, after=7, line=1.08)
        add_text(p2, title, size=10.8, color=INK, bold=True)
        p3 = cell.add_paragraph()
        style_paragraph(p3, after=0, line=1.16)
        add_text(p3, body, size=8.6, color=GRAPHITE)
    blank_line(doc, 12)


def add_overview_page(doc):
    doc.add_page_break()
    add_topline(
        doc,
        "02",
        "Executive Overview",
        "A controlled pilot should prove demand, workflow, commercial return and customer delivery before broader rollout.",
    )

    table = doc.add_table(rows=1, cols=2)
    set_table_fixed_width(table, [3000, 6994])
    clear_table_borders(table)
    side, body = table.cell(0, 0), table.cell(0, 1)
    set_cell_shading(side, SOFT_BLUE)
    set_cell_margins(side, top=180, bottom=180, start=170, end=170)
    set_cell_margins(body, top=0, bottom=0, start=240, end=0)
    set_cell_border(
        side,
        top={"val": "single", "sz": "4", "color": LINE},
        bottom={"val": "single", "sz": "4", "color": LINE},
        left={"val": "single", "sz": "4", "color": LINE},
        right={"val": "single", "sz": "4", "color": LINE},
    )

    p = side.paragraphs[0]
    style_paragraph(p, after=10, line=1.0)
    add_text(p, "DISCUSSION PAPER", size=7.8, color=BLUE, bold=True)
    p = side.add_paragraph()
    style_paragraph(p, after=12, line=1.08)
    add_text(p, "Collaborative, not prescriptive.", size=15.5, color=INK, bold=True)
    p = side.add_paragraph()
    style_paragraph(p, after=18, line=1.2)
    add_text(
        p,
        "The tone should invite Afgri into the programme design while still giving them a practical path to start.",
        size=9.4,
        color=GRAPHITE,
    )
    for label, value in [
        ("Pilot scope", "Warehouses, ground mounts, carports"),
        ("Marketing lead", "Afgri customer and branch channels"),
        ("Commercial return", "5% proposed return plus tier incentives"),
    ]:
        p = side.add_paragraph()
        style_paragraph(p, after=2, line=1.0)
        add_text(p, label.upper(), size=7.2, color=BLUE, bold=True)
        p2 = side.add_paragraph()
        style_paragraph(p2, after=9, line=1.12)
        add_text(p2, value, size=8.8, color=INK)

    p = body.paragraphs[0]
    style_paragraph(p, after=8, line=1.25)
    add_text(
        p,
        "Smart Steel can offer Afgri a practical infrastructure product range, supported by Atlas for quoting, workflow visibility and pilot control. Afgri brings an established agricultural customer base, financing relationships and a trusted branch network.",
        size=10.2,
        color=INK,
    )
    p = body.add_paragraph()
    style_paragraph(p, after=12, line=1.25)
    add_text(
        p,
        "The proposed first step is not a full national rollout. It is a measured pilot that allows both parties to validate the commercial model, installation options, marketing response and internal handoffs before scaling.",
        size=10.2,
        color=INK,
    )
    add_callout(
        doc,
        "Pilot start mechanism",
        "The document can include a signed initial offer or heads of terms so the programme can begin while the broader final agreement is completed.",
    )


def add_commercial_page(doc):
    doc.add_page_break()
    add_topline(
        doc,
        "03",
        "Commercial Framework",
        "A proposed pilot return that protects Smart Steel pricing authority while giving Afgri a clear reason to activate demand.",
    )

    add_signal_cards(
        doc,
        [
            ("Base", "5% pilot return", "A proposed commercial return on qualifying pilot sales, subject to final modelling and agreement.", PALE_BLUE),
            ("Growth", "Tiered incentive", "Additional volume-based upside can be linked to agreed quarterly or pilot-period thresholds.", "FFFFFF"),
            ("Control", "Protected pricing", "Smart Steel retains technical pricing authority while Afgri receives a structured route to market.", SOFT_BLUE),
        ],
    )

    table = doc.add_table(rows=4, cols=4)
    set_table_fixed_width(table, [1900, 3310, 2810, 1974])
    clear_table_borders(table)
    headers = ["Tier", "Trigger", "Commercial Return", "Purpose"]
    for c, text in enumerate(headers):
        cell = table.cell(0, c)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=115, bottom=115, start=110, end=110)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.0)
        add_text(p, text, size=8.5, color=RGBColor(255, 255, 255), bold=True)
    rows = [
        ("Pilot", "Qualifying accepted orders", "5% proposed return", "Launch activation"),
        ("Volume", "Threshold to be modelled", "Incremental incentive", "Reward momentum"),
        ("Strategic", "Broader rollout approval", "Final negotiated model", "Scale responsibly"),
    ]
    for r, row in enumerate(rows, start=1):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            set_cell_shading(cell, "FFFFFF" if r != 2 else PALE_GRAY)
            set_cell_margins(cell, top=125, bottom=125, start=110, end=110)
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": "4", "color": LINE},
                left={"val": "single", "sz": "4", "color": LINE},
                right={"val": "single", "sz": "4", "color": LINE},
            )
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.12)
            add_text(p, text, size=8.7, color=BLUE if c == 2 else INK, bold=c in (0, 2))
    blank_line(doc, 10)
    add_callout(
        doc,
        "Commercial wording",
        "Position the 5% as a proposed pilot commercial return, subject to final modelling, qualification rules and written agreement between the parties.",
    )


def add_marketing_page(doc):
    doc.add_page_break()
    add_topline(
        doc,
        "04",
        "Market Activation",
        "Afgri should lead the customer-facing campaign, with Smart Steel supporting the product and technical layer.",
    )
    add_signal_cards(
        doc,
        [
            ("Lead", "Afgri campaign owner", "Afgri drives branch, customer, finance and digital activation as the trusted market channel.", PALE_BLUE),
            ("Support", "Smart Steel content", "Smart Steel supplies product copy, technical visuals, quote support and proof points for the campaign.", "FFFFFF"),
            ("Measure", "Atlas visibility", "Lead source, quote status and conversion can be tracked through the pilot workflow.", SOFT_BLUE),
        ],
    )
    add_heading(doc, "Campaign posture", 2)
    table = doc.add_table(rows=3, cols=3)
    set_table_fixed_width(table, [1900, 2794, 5300])
    clear_table_borders(table)
    rows = [
        ("01", "Afgri leads", "Afgri is positioned as the primary customer-facing channel for the pilot campaign."),
        ("02", "Smart Steel supports", "Smart Steel remains visible as the product and manufacturing partner behind Atlas-enabled infrastructure systems."),
        ("03", "Assets stay practical", "Product pages, branch material, finance-linked offers, lead forms and simple campaign reporting support conversion."),
    ]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            set_cell_shading(cell, PALE_BLUE if c == 0 else "FFFFFF")
            set_cell_margins(cell, top=115, bottom=115, start=110, end=110)
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": "4", "color": LINE},
                left={"val": "single", "sz": "4", "color": LINE},
                right={"val": "single", "sz": "4", "color": LINE},
            )
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c == 0 else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.12)
            add_text(p, text, size=8.8, color=BLUE if c == 0 else INK, bold=c < 2)
    blank_line(doc, 12)
    add_callout(
        doc,
        "Why Afgri leads",
        "Afgri has the broader agricultural reach, stronger customer trust and more relevant customer data. Smart Steel marketing can support conversion, but Afgri should lead demand creation.",
    )


def add_workflow_page(doc):
    doc.add_page_break()
    add_topline(
        doc,
        "05",
        "Pilot Workflow and Governance",
        "A simple operating rhythm helps the pilot start quickly without losing control of pricing, scope changes, installation responsibility or customer expectations.",
    )

    table = doc.add_table(rows=2, cols=5)
    set_table_fixed_width(table, [1998, 1998, 1998, 1998, 2002])
    clear_table_borders(table)
    steps = [
        ("01", "Lead"),
        ("02", "Technical Review"),
        ("03", "Dealer Quote"),
        ("04", "Customer Proposal"),
        ("05", "Order and Delivery"),
    ]
    for c, (num, label) in enumerate(steps):
        top = table.cell(0, c)
        bottom = table.cell(1, c)
        set_cell_shading(top, NAVY if c == 0 else PALE_BLUE)
        set_cell_shading(bottom, "FFFFFF")
        for cell in [top, bottom]:
            set_cell_margins(cell, top=115, bottom=115, start=95, end=95)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": LINE},
                bottom={"val": "single", "sz": "4", "color": LINE},
                left={"val": "single", "sz": "4", "color": LINE},
                right={"val": "single", "sz": "4", "color": LINE},
            )
        p = top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=0, line=1.0)
        step_color = RGBColor(255, 255, 255) if c == 0 else BLUE
        add_text(p, num, size=10.5, color=step_color, bold=True)
        p = bottom.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=0, line=1.08)
        add_text(p, label, size=8.4, color=INK, bold=True)
    blank_line(doc, 16)

    table = doc.add_table(rows=4, cols=3)
    set_table_fixed_width(table, [2500, 3994, 3500])
    clear_table_borders(table)
    rows = [
        ("Governance", "Weekly pilot review", "Pipeline, blockers, pricing exceptions and installation status."),
        ("Scope changes", "Written variation route", "Any customer-specific changes move through technical review before revised pricing."),
        ("Cancellation", "Clear cost recovery", "Protect work already completed, ordered materials and engineering time."),
        ("Installation", "Optional included scope", "Installation can be quoted where the customer wants a complete delivered solution."),
    ]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            set_cell_shading(cell, PALE_GRAY if c == 0 else "FFFFFF")
            set_cell_margins(cell, top=120, bottom=120, start=120, end=120)
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": "4", "color": LINE},
                left={"val": "single", "sz": "4", "color": LINE},
                right={"val": "single", "sz": "4", "color": LINE},
            )
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.12)
            add_text(p, text, size=8.7, color=BLUE if c == 0 else INK, bold=c == 0)


def add_component_page(doc):
    doc.add_page_break()
    add_heading(doc, "03 Component System", 1)
    add_body(
        doc,
        "This page demonstrates the reusable body components for the Afgri-facing paper: concise paragraphs, strong callouts, compact bullets and fixed-width tables."
    )
    add_callout(
        doc,
        "Pilot principle",
        "The programme should be able to begin from a signed pilot heads of terms while the broader legal and operational documents are finalised."
    )
    add_heading(doc, "Reusable Signals", 2)
    for item in [
        "Afgri leads customer-facing marketing and branch activation.",
        "Smart Steel supports technical accuracy, Atlas workflow and product documentation.",
        "The pilot validates commercial, financing, installation and dashboard assumptions before broader rollout.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Alignment Matrix", 2)
    add_sample_table(doc)
    blank_line(doc, 16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run().add_picture(str(ATLAS_LOGO), width=Inches(1.55))


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_doc(doc)
    set_running_furniture(doc.sections[0])
    add_section_opener(
        doc,
        "01",
        "Partnership\nRationale",
        "Afgri reach + Smart Steel infrastructure systems\n+ Atlas workflow.",
        [
            ("Market reach", "Afgri leads the customer-facing campaign through branch, finance and customer channels.", PALE_BLUE),
            ("Technical depth", "Smart Steel supports product configuration, engineering review and manufacturing readiness.", "FFFFFF"),
            ("Pilot control", "The first phase validates demand, process, installation and commercial assumptions.", SOFT_BLUE),
        ],
        start_new_page=False,
    )
    add_overview_page(doc)
    add_commercial_page(doc)
    add_marketing_page(doc)
    add_workflow_page(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
